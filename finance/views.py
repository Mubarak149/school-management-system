import json
import openpyxl
import io
import os
import weasyprint
from datetime import timedelta
from django.conf import settings
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse, HttpResponse
from django.template.loader import render_to_string
from django.utils import timezone
from django.db import transaction, DatabaseError
from django.db.models import Sum, Count, Q
from django.core.paginator import Paginator
from .forms import (InvoiceForm, PaymentForm, InvoiceItemForm,
                    FeeCategoryForm, FeeStructureForm)
from .models import (SchoolInvoice, InvoiceItem, 
                     Payment, FeeStructure, FeeCategory)
from student.models import Student
from school_admin.models import SiteSetting



def finance_dashboard(request):
    # Aggregate stats
    invoice_status_counts = SchoolInvoice.objects.aggregate(
        unpaid=Count("id", filter=Q(status="unpaid")),
        partial=Count("id", filter=Q(status="partial")),
        paid=Count("id", filter=Q(status="paid")),
    )

    total_students = Student.objects.count()
    total_collected = Payment.objects.aggregate(total=Sum("amount"))["total"] or 0

    # --- PAGINATION FOR INVOICES ---
    invoice_list = SchoolInvoice.objects.all().only(
        "id", "invoice_number", "status", "due_date", "student_id"
    ).select_related("student").order_by("-id")
    invoice_paginator = Paginator(invoice_list, 5)
    invoices = invoice_paginator.get_page(request.GET.get("invoice_page"))

    # --- PAGINATION FOR PAYMENTS ---
    payment_list = Payment.objects.all().only(
        "id", "amount", "date_paid", "payment_method", "invoice_id"
    ).select_related("invoice").order_by("-id")
    payment_paginator = Paginator(payment_list, 5)
    payments = payment_paginator.get_page(request.GET.get("payment_page"))

    # --- PAGINATION FOR FEE STRUCTURE ---
    fee_structure_list = FeeStructure.objects.all().only(
        "id", "amount", "term", "school_class_id", "category_id"
    ).select_related("school_class", "category").order_by("-id")
    fee_structure_paginator = Paginator(fee_structure_list, 5)
    fee_structure = fee_structure_paginator.get_page(request.GET.get("structure_page"))

    # --- UPCOMING DUE & OVERDUE INVOICES ---
    today = timezone.now().date()
    upcoming_due = SchoolInvoice.objects.filter(
        due_date__gte=today,
        due_date__lte=today + timedelta(days=3),
        status__in=["unpaid", "partial"]
    ).select_related("student")

    overdue = SchoolInvoice.objects.filter(
        due_date__lt=today,
        status__in=["unpaid", "partial"]
    ).select_related("student")

    # Forms + Context
    category = FeeCategory.objects.all().only("id", "name")
    context = {
        "students": total_students,
        "invoices": invoices,
        "payments": payments,
        "structures": fee_structure,
        "invoice_form": InvoiceForm(),
        "categories": category,
        "category_form": FeeCategoryForm(),
        "fee_structure_form": FeeStructureForm(),
        "total_students": total_students,
        "students_paid": invoice_status_counts["paid"],
        "students_unpaid": invoice_status_counts["unpaid"] + invoice_status_counts["partial"],
        "total_collected": total_collected,
        "unpaid_invoices": invoice_status_counts["unpaid"],
        "partial_invoices": invoice_status_counts["partial"],
        "paid_invoices": invoice_status_counts["paid"],
        "upcoming_due": upcoming_due,
        "overdue": overdue,
    }
    return render(request, "finance/finance_dashboard.html", context)


#category view
def category_row(request, pk):
    category = get_object_or_404(FeeCategory, pk=pk)
    html = render_to_string("finance/partials/category_row.html", {"cat": category})
    return HttpResponse(html)

def create_category(request):
    if request.method == "POST":
        form = FeeCategoryForm(request.POST)
        if form.is_valid():
            category = form.save()
            html = render_to_string("finance/partials/category_row.html", {"cat": category}, request=request)
            return HttpResponse(html)  # HTMX swaps this into table
        return JsonResponse({"error": "Invalid data"}, status=400)

def edit_category(request, pk):
    category = get_object_or_404(FeeCategory, pk=pk)
    if request.method == "GET":
        form = FeeCategoryForm(instance=category)
        return render(request, "finance/partials/category_edit_form.html", {"form": form, "cat": category})

    elif request.method == "POST":
        form = FeeCategoryForm(request.POST, instance=category)
        if form.is_valid():
            category = form.save()
            html = render_to_string("finance/partials/category_row.html", {"cat": category})
            return HttpResponse(html)
        return JsonResponse({"error": "Invalid data"}, status=400)

def delete_category(request, pk):
    category = get_object_or_404(FeeCategory, pk=pk)
    category.delete()
    return HttpResponse("")  # removes row

def fee_category_dropdown(request):
    categories = FeeCategory.objects.all()
    return render(request, "finance/partials/fee_category_dropdown.html", {"categories": categories})


# Fee structure View
def fee_structures_view(request):
    structures = FeeStructure.objects.select_related("school_class", "session", "category").all()
    form = FeeStructureForm()
    return render(request, "finance/fee_structures.html", {
        "structures": structures,
        "form": form,
    })

def create_fee_structure(request):
    if request.method == "POST":
        form = FeeStructureForm(request.POST)
        if form.is_valid():
            structure = form.save()
            
            # Render the row for updating the table
            html = render_to_string(
                "finance/partials/fee_structure_row.html",
                {"fs": structure},
                request=request
            )

            # ✅ Success message
            message = f"✅ Fee structure '{structure}' created successfully."

            response = HttpResponse(html)
            response["HX-Trigger"] = json.dumps({
                "showMessage": {"message": message, "type": "success"}
            })
            return response

        # ❌ Invalid form
        response = HttpResponse("Invalid data", status=400)
        response["HX-Trigger"] = json.dumps({
            "showMessage": {"message": "❌ Invalid fee structure data.", "type": "error"}
        })
        return response

def edit_fee_structure(request, pk):
    fs = get_object_or_404(FeeStructure, pk=pk)
    if request.method == "POST":
        form = FeeStructureForm(request.POST, instance=fs)
        if form.is_valid():
            fs = form.save()
            html = render_to_string("finance/partials/fee_structure_row.html", {"fs": fs},request=request)
            return HttpResponse(html)
        return HttpResponse("Invalid data", status=400)

    form = FeeStructureForm(instance=fs)
    html = render_to_string("finance/partials/fee_structure_edit.html", {"form": form, "fs": fs},request=request)
    return HttpResponse(html)

def cancel_edit_fee_structure(request, pk):
    fs = get_object_or_404(FeeStructure, pk=pk)
    html = render_to_string("finance/partials/fee_structure_row.html", {"fs": fs},request=request)
    return HttpResponse(html)

def delete_fee_structure(request, pk):
    fs = get_object_or_404(FeeStructure, pk=pk)
    fs.delete()
    return HttpResponse("")  # HTMX will remove the row

def fee_structure_page(request, page):
    structures = FeeStructure.objects.all().order_by("school_class")
    paginator = Paginator(structures, 5)
    page_obj = paginator.get_page(page)
    return render(request, "finance/partials/fee_structure_table.html", {"structures": page_obj})

def fee_structure_invoices(request, fs_id):
    fs = get_object_or_404(FeeStructure, id=fs_id)

    # Get invoices related to this class/session/term
    invoices = SchoolInvoice.objects.filter(
        student__myclasses__student_class=fs.school_class,
        session=fs.session,
        term=fs.term
    ).select_related("student", "session").prefetch_related("items", "payments")

    return render(request, "finance/fee_structure_invoices.html", {
        "fs": fs,
        "invoices": invoices,
    })

# Create Payment
def create_payment(request):
    if request.method == "POST":
        form = PaymentForm(request.POST)
        if form.is_valid():
            form.save()
            payments = Payment.objects.all().order_by("-date_paid")
            return render(request, "finance/partials/payments_table_rows.html", {
                "payments": payments
            })
    return HttpResponse(status=405)

def invoice_payment(request, invoice_id):
    invoice = get_object_or_404(SchoolInvoice, id=invoice_id)

    if request.method == "POST":
        form = PaymentForm(request.POST)
        if form.is_valid():
            payment = form.save(commit=False)
            payment.invoice = invoice   # ✅ attach manually
            payment.date_paid = timezone.now()
            payment.save()
            invoice.update_status()
            return redirect("invoice_detail", pk=invoice.id)
    else:
        form = PaymentForm()

    return render(request, "finance/invoice_payment.html", {
        "invoice": invoice,
        "form": form
    })

def payment_page(request, page):
    payments = Payment.objects.all()
    paginator = Paginator(payments, 20)
    page_obj = paginator.get_page(page)
    return render(request, "finance/partials/payments_table.html", {"payments": page_obj})


def generate_receipt(request, invoice_id):
    invoice = get_object_or_404(SchoolInvoice, id=invoice_id)
    payments = invoice.payments.all()

    # Render HTML
    html_string = render_to_string("finance/receipt.html", {
        "invoice": invoice,
        "payments": payments,
    })

    # Convert to PDF
    response = HttpResponse(content_type="application/pdf")
    response['Content-Disposition'] = f'inline; filename=receipt_{invoice.invoice_number}.pdf'
    weasyprint.HTML(string=html_string).write_pdf(response)

    return response


#invoice view
def invoice_page(request, page):
    invoices = SchoolInvoice.objects.all().order_by("-due_date")
    paginator = Paginator(invoices, 20)
    page_obj = paginator.get_page(page)
    return render(request, "finance/partials/invoice_list.html", {"invoices": page_obj})

def invoice_detail(request, pk):
    invoice = get_object_or_404(SchoolInvoice, pk=pk)
    return render(request, "finance/invoice_detail.html", {"invoice": invoice})

def invoice_detail_pdf(request, pk):
    invoice = get_object_or_404(SchoolInvoice, pk=pk)
    site = SiteSetting.objects.first()  # assuming you always have 1 settings record

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    # --- Header with Logo & School Name ---
    if site and site.logo:
        logo_path = os.path.join(settings.MEDIA_ROOT, str(site.logo))
        if os.path.exists(logo_path):
            img = Image(logo_path, width=80, height=80)  # adjust size
            header_data = [
                [img, Paragraph(f"<b>{site.school_name}</b>", styles["Title"])]
            ]
            header_table = Table(header_data, colWidths=[100, 400])
            header_table.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            elements.append(header_table)
            elements.append(Spacer(1, 20))
    else:
        # fallback if no logo
        elements.append(Paragraph(f"<b>{site.school_name if site else 'School Name'}</b>", styles["Title"]))
        elements.append(Spacer(1, 20))

    # Invoice Info
    elements.append(Paragraph(f"<b>Invoice:</b> #{invoice.invoice_number}", styles["Heading2"]))
    elements.append(Paragraph(f"<b>Student:</b> {invoice.student}", styles["Normal"]))
    elements.append(Paragraph(f"<b>Status:</b> {invoice.get_status_display()}", styles["Normal"]))
    elements.append(Paragraph(f"<b>Issue Date:</b> {invoice.issue_date}", styles["Normal"]))
    elements.append(Paragraph(f"<b>Due Date:</b> {invoice.due_date}", styles["Normal"]))
    elements.append(Spacer(1, 15))

    # Items Table
    data = [["Category", "Description", "Amount (₦)"]]
    for item in invoice.items.all():
        data.append([str(item.category), item.description or "-", f"{item.amount:,.2f}"])

    table = Table(data, colWidths=[150, 250, 100])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), "#333333"),
        ("TEXTCOLOR", (0, 0), (-1, 0), "#FFFFFF"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
        ("BACKGROUND", (0, 1), (-1, -1), "#F9F9F9"),
        ("GRID", (0, 0), (-1, -1), 0.5, "#000000"),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 20))

    # Totals
    totals_data = [
        ["", "Total", f"₦{invoice.total_amount():,.2f}"],
        ["", "Paid", f"₦{invoice.amount_paid():,.2f}"],
        ["", "Balance", f"₦{invoice.balance():,.2f}"],
    ]
    totals_table = Table(totals_data, colWidths=[250, 150, 100])
    totals_table.setStyle(TableStyle([
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("FONTNAME", (1, 0), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("GRID", (1, 0), (-1, -1), 0.5, "#000000"),
        ("BACKGROUND", (1, 0), (-1, -1), "#EFEFEF"),
    ]))
    elements.append(totals_table)

    # Footer with contact info if available
    if site and site.address:
        elements.append(Spacer(1, 30))
        elements.append(Paragraph(f"<b>Address:</b> {site.address}", styles["Normal"]))
    if site and site.contact_phone:
        elements.append(Paragraph(f"<b>Phone:</b> {site.contact_phone}", styles["Normal"]))
    if site and site.contact_email:
        elements.append(Paragraph(f"<b>Email:</b> {site.contact_email}", styles["Normal"]))

    # Build PDF
    doc.build(elements)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename=invoice_{invoice.invoice_number}.pdf'
    return response

def invoice_detail_word(request, pk):
    invoice = get_object_or_404(SchoolInvoice, pk=pk)
    site = SiteSetting.objects.first()

    doc = Document()

    # --- Header with Logo & School Name ---
    if site:
        table = doc.add_table(rows=1, cols=2)
        row = table.rows[0]

        if site.logo:
            logo_path = os.path.join(settings.MEDIA_ROOT, str(site.logo))
            if os.path.exists(logo_path):
                cell_logo = row.cells[0]
                paragraph = cell_logo.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Inches(1), height=Inches(1))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            row.cells[0].text = ""

        # School Name
        cell_school = row.cells[1]
        p = cell_school.paragraphs[0]
        run = p.add_run(site.school_name)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")  # space

    # --- Invoice Info ---
    doc.add_heading(f"Invoice #{invoice.invoice_number}", level=1)
    doc.add_paragraph(f"Student: {invoice.student}")
    doc.add_paragraph(f"Status: {invoice.get_status_display()}")
    doc.add_paragraph(f"Issue Date: {invoice.issue_date}")
    doc.add_paragraph(f"Due Date: {invoice.due_date}")

    doc.add_paragraph("")  # space

    # --- Items Table ---
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"  # Word built-in style
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Category"
    hdr_cells[1].text = "Description"
    hdr_cells[2].text = "Amount (₦)"

    total = 0
    for item in invoice.items.all():
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.category)
        row_cells[1].text = item.description or "-"
        row_cells[2].text = f"{item.amount:,.2f}"
        total += item.amount

    # --- Totals Row ---
    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = "TOTAL"
    total_row[2].text = f"₦{total:,.2f}"

    for cell in total_row:
        for paragraph in cell.paragraphs:
            if paragraph.runs:
                run = paragraph.runs[0]
                run.bold = True
                run.font.size = Pt(11)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("")  # space

    # --- Footer with Contact Info ---
    if site:
        if site.address:
            doc.add_paragraph(f"Address: {site.address}")
        if site.contact_phone:
            doc.add_paragraph(f"Phone: {site.contact_phone}")
        if site.contact_email:
            doc.add_paragraph(f"Email: {site.contact_email}")

    # Save & Response
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="invoice_{invoice.invoice_number}.docx"'
    return response

def send_invoice(request, fs_id):
    fs = get_object_or_404(FeeStructure, id=fs_id)

    students = Student.objects.filter(
        myclasses__student_class=fs.school_class,
        myclasses__current_class=True
    )
    fee_structures = FeeStructure.objects.filter(
        school_class=fs.school_class,
        session=fs.session,
        term=fs.term
    ).select_related("category")

    invoices_created = []
    updated_invoices = []  # 🆕 track invoices where new items were added
    errors = []

    try:
        with transaction.atomic():
            for student in students:
                existing_invoice = SchoolInvoice.objects.filter(
                    student=student, session=fs.session, term=fs.term
                ).first()

                if existing_invoice:
                    missing_items = False
                    for struct in fee_structures:
                        item, created = InvoiceItem.objects.get_or_create(
                            invoice=existing_invoice,
                            category=struct.category,
                            defaults={
                                "description": struct.category.name,
                                "amount": struct.amount,
                            }
                        )
                        if created:
                            missing_items = True

                    if missing_items:
                        updated_invoices.append(existing_invoice.invoice_number)
                    else:
                        errors.append(f"Invoice already exists for {student}")
                    continue

                # --- Create new invoice ---
                now = timezone.now()
                invoice_number = f"INV-{now.year}-{int(now.timestamp() * 1_000_000)}"

                invoice = SchoolInvoice.objects.create(
                    student=student,
                    invoice_number=invoice_number,
                    session=fs.session,
                    term=fs.term,
                    issue_date=timezone.now(),
                    due_date=timezone.now() + timezone.timedelta(days=14)
                )

                items = [
                    InvoiceItem(
                        invoice=invoice,
                        category=struct.category,
                        description=struct.category.name,
                        amount=struct.amount
                    )
                    for struct in fee_structures
                ]
                InvoiceItem.objects.bulk_create(items)
                invoices_created.append(invoice_number)

        # ✅ Success / Info / Warning messages
        if invoices_created:
            message = f"✅ {len(invoices_created)} invoice(s) created successfully."
        if updated_invoices:
            message = f"➕ Added new items to {len(updated_invoices)} existing invoice(s)."
        if errors and not (invoices_created or updated_invoices):
            message = "⚠️ All invoices already exist. No changes made."

        response = HttpResponse("")
        response["HX-Trigger"] = json.dumps({
            "showMessage": {"message": message, "type": "success"}
        })
        return response

    except Exception as e:
        response = HttpResponse("", status=500)
        response["HX-Trigger"] = json.dumps({
            "showMessage": {"message": f"❌ Error: {str(e)}", "type": "danger"}
        })
        return response


def export_invoices_pdf(request, fs_id):
    fs = get_object_or_404(FeeStructure, id=fs_id)
    invoices = SchoolInvoice.objects.filter(
        student__myclasses__student_class=fs.school_class,
        session=fs.session,
        term=fs.term
    ).select_related("student")

    # Create PDF response
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="invoices_{fs.school_class}_{fs.session}_{fs.term}.pdf"'

    p = canvas.Canvas(response, pagesize=A4)
    width, height = A4

    # Title
    p.setFont("Helvetica-Bold", 14)
    p.drawString(1 * inch, height - 1 * inch, f"Invoices for {fs.school_class} - {fs.session} ({fs.term})")

    y = height - 1.5 * inch
    p.setFont("Helvetica", 10)

    # Table Header
    p.drawString(1 * inch, y, "Invoice #")
    p.drawString(2.8 * inch, y, "Student")
    p.drawString(4.8 * inch, y, "Total")
    p.drawString(5.6 * inch, y, "Balance")
    p.drawString(6.4 * inch, y, "Status")
    y -= 0.3 * inch

    # Rows
    for invoice in invoices:
        if y < 1 * inch:  # New page if space is low
            p.showPage()
            y = height - 1 * inch
            p.setFont("Helvetica", 10)

        p.drawString(1 * inch, y, invoice.invoice_number)
        p.drawString(2.8 * inch, y, f"{invoice.student}")
        p.drawString(4.8 * inch, y, f"₦{invoice.total_amount():,.2f}")
        p.drawString(5.6 * inch, y, f"₦{invoice.balance():,.2f}")
        p.drawString(6.4 * inch, y, invoice.get_status_display())
        y -= 0.25 * inch

    p.showPage()
    p.save()
    return response

def export_invoices_word(request, fs_id):
    fs = get_object_or_404(FeeStructure, id=fs_id)
    invoices = SchoolInvoice.objects.filter(
        student__myclasses__student_class=fs.school_class,
        session=fs.session,
        term=fs.term
    )

    document = Document()
    document.add_heading(f"Invoices for {fs.school_class} - {fs.term} ({fs.session})", 0)

    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Invoice No"
    hdr_cells[1].text = "Student"
    hdr_cells[2].text = "Balance"
    hdr_cells[3].text = "Status"
    hdr_cells[4].text = "Session/Term"

    for inv in invoices:
        row = table.add_row().cells
        row[0].text = str(inv.invoice_number)
        row[1].text = str(inv.student)
        row[2].text = f"₦{inv.balance():.2f}"
        row[3].text = inv.get_status_display()
        row[4].text = f"{inv.session} / {inv.term}"

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="invoices_{fs.id}.docx"'
    document.save(response)
    return response

def export_invoices_excel(request, fs_id):
    fs = get_object_or_404(FeeStructure, id=fs_id)
    invoices = SchoolInvoice.objects.filter(
        student__myclasses__student_class=fs.school_class,
        session=fs.session,
        term=fs.term
    )

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Invoices"

    ws.append(["Invoice No", "Student", "Balance", "Status", "Session", "Term"])

    for inv in invoices:
        ws.append([
            inv.invoice_number,
            str(inv.student),
            float(inv.balance()),
            inv.get_status_display(),
            str(inv.session),
            str(inv.term),
        ])

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response["Content-Disposition"] = f'attachment; filename="invoices_{fs.id}.xlsx"'
    wb.save(response)
    return response

